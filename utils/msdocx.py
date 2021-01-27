from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

class Docx:

    MNAME = "docx"

    def __init__(self, doc_name="test.docx", user_folder=None):
        self.document = Document()
        self.doc_name = doc_name
        self.user_folder = user_folder

    def save_as(self):
        self.document.save(self.doc_name)

    def insert_images(self):
        images = os.listdir(self.user_folder)
        size = len(images)
        for i, img in enumerate(images):
            print(f"adding image #{i + 1} out of {size}")
            self.document.add_picture(self.user_folder + "/" + img, width=Inches(5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def change_to_landscape(self):
        section = self.document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE

    def change_to_portrait(self):
        section = self.document.sections[-1]
        section.orientation = WD_ORIENT.PORTRAIT

    def add_single_page(self):
        self.document.add_page_break()

    


    
        
    
